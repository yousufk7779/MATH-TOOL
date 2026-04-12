from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE_TYPE

doc = Document(r'D:\All NCERT SOLUTIONS\source_docs\science\physics\chapter3.docx')

found_q5 = False
for i, para in enumerate(doc.paragraphs):
    if "Use the data in Table" in para.text and ("iron and mercury" in para.text or "better conductor" in para.text):
        print(f"Found Question 5 at paragraph {i}: {para.text}")
        found_q5 = True
        # Look for images in this and subsequent paragraphs
        for j in range(i, i + 10):
            if j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                for run in p.runs:
                    if run.element.xpath('.//a:blip'):
                        print(f"Found image reference in paragraph {j}")
                        # I'll need more logic to map blip to image name, but let's see if this prints anything.
